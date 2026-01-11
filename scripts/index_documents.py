#!/usr/bin/env python3
"""
Document Indexer for Enigma
Creates semantic embeddings of all documents for fast retrieval.

Usage:
    python3 scripts/index_documents.py
"""

import os
from pathlib import Path
import chromadb
from chromadb.config import Settings
from docx import Document
import pymupdf  # PyMuPDF for PDF reading

# Enigma paths
ENIGMA_ROOT = Path(__file__).parent.parent
CHROMA_DIR = ENIGMA_ROOT / "data" / "chromadb"
GROUND_TRUTH_DIR = ENIGMA_ROOT / "99_GroundTruth"
SYSTEMS_DIR = ENIGMA_ROOT / "02_Systems"
PROJECTS_DIR = ENIGMA_ROOT / "03_Projects"
MEMORY_DIR = ENIGMA_ROOT / "04_Memory"

def read_text_file(file_path):
    """Read plain text or markdown file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def read_word_file(file_path):
    """Read Word document (.docx)."""
    try:
        doc = Document(str(file_path))
        # Extract all paragraphs
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)

        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)

        content = '\n\n'.join(text)
        if not content.strip():
            raise Exception("Document appears to be empty")
        return content
    except Exception as e:
        raise Exception(f"Error reading Word file: {str(e)}")

def read_pdf_file(file_path):
    """Read PDF file."""
    try:
        doc = pymupdf.open(file_path)
        text = []
        for page in doc:
            text.append(page.get_text())
        doc.close()
        return '\n\n'.join(text)
    except Exception as e:
        raise Exception(f"Error reading PDF file: {e}")

def chunk_document(content, filename, chunk_size=1000, overlap=200):
    """Split document into overlapping chunks."""
    chunks = []
    words = content.split()

    for i in range(0, len(words), chunk_size - overlap):
        chunk = ' '.join(words[i:i + chunk_size])
        chunks.append({
            'content': chunk,
            'filename': filename,
            'chunk_id': f"{filename}_chunk_{len(chunks)}"
        })

    return chunks

def load_documents(directory, layer):
    """Load all supported document types from a directory."""
    documents = []

    if not directory.exists():
        print(f"⚠️  Directory not found: {directory}")
        return documents

    # Supported file extensions and their readers
    file_readers = {
        '.md': read_text_file,
        '.txt': read_text_file,
        '.docx': read_word_file,
        '.doc': read_word_file,  # Old Word format (best effort)
        '.pdf': read_pdf_file
    }

    # Find all supported files
    for file_path in directory.rglob("*"):
        if not file_path.is_file():
            continue

        file_ext = file_path.suffix.lower()

        if file_ext not in file_readers:
            continue

        try:
            # Read file using appropriate reader
            reader = file_readers[file_ext]
            content = reader(file_path)

            # Skip if empty
            if not content.strip():
                continue

            # Chunk the document
            chunks = chunk_document(content, file_path.name)

            for chunk in chunks:
                documents.append({
                    'id': chunk['chunk_id'],
                    'content': chunk['content'],
                    'metadata': {
                        'filename': chunk['filename'],
                        'filetype': file_ext,
                        'layer': layer,
                        'path': str(file_path.relative_to(ENIGMA_ROOT))
                    }
                })

            print(f"  ✅ {file_path.name} ({file_ext}): {len(chunks)} chunks")

        except Exception as e:
            print(f"  ⚠️  Error reading {file_path.name}: {e}")

    return documents

def create_index():
    """Create or update the document index."""

    print("="*80)
    print("ENIGMA DOCUMENT INDEXER")
    print("="*80)
    print()

    # Create data directory if needed
    CHROMA_DIR.parent.mkdir(parents=True, exist_ok=True)

    # Initialize ChromaDB
    print("🔧 Initializing ChromaDB...")
    client = chromadb.PersistentClient(path=str(CHROMA_DIR))

    # Delete existing collection if it exists
    try:
        client.delete_collection("enigma_knowledge")
        print("  ✅ Cleared existing index")
    except:
        pass

    # Create new collection
    collection = client.create_collection(
        name="enigma_knowledge",
        metadata={"description": "Enigma knowledge base"}
    )
    print("  ✅ Created new index\n")

    # Load documents from each layer
    all_documents = []

    print("📚 Loading Layer 1: Ground Truth...")
    print("   Supported: .md, .txt, .doc, .docx, .pdf")
    all_documents.extend(load_documents(GROUND_TRUTH_DIR, "Layer 1: Ground Truth"))

    print("\n📚 Loading Layer 2: Systems...")
    print("   Supported: .md, .txt, .doc, .docx, .pdf")
    all_documents.extend(load_documents(SYSTEMS_DIR, "Layer 2: Structured Reality"))

    print("\n📚 Loading Layer 2: Memory...")
    print("   Supported: .md, .txt, .doc, .docx, .pdf")
    all_documents.extend(load_documents(MEMORY_DIR, "Layer 2: Structured Reality"))

    print("\n📚 Loading Layer 3: Projects...")
    print("   Supported: .md, .txt, .doc, .docx, .pdf")
    all_documents.extend(load_documents(PROJECTS_DIR, "Layer 3: Working Knowledge"))

    print()
    print("="*80)
    print(f"📊 INDEXING SUMMARY")
    print("="*80)
    print(f"Total chunks to index: {len(all_documents)}")

    if not all_documents:
        print("\n❌ No documents found to index!")
        return

    # Add to ChromaDB
    print("\n🔍 Creating embeddings and indexing...")

    ids = [doc['id'] for doc in all_documents]
    contents = [doc['content'] for doc in all_documents]
    metadatas = [doc['metadata'] for doc in all_documents]

    collection.add(
        ids=ids,
        documents=contents,
        metadatas=metadatas
    )

    print("  ✅ Indexing complete!")

    # Show summary by layer
    print("\n📊 Documents by layer:")
    layer_counts = {}
    for doc in all_documents:
        layer = doc['metadata']['layer']
        layer_counts[layer] = layer_counts.get(layer, 0) + 1

    for layer, count in sorted(layer_counts.items()):
        print(f"  {layer}: {count} chunks")

    # Show summary by file type
    print("\n📊 Documents by file type:")
    filetype_counts = {}
    for doc in all_documents:
        filetype = doc['metadata'].get('filetype', 'unknown')
        filetype_counts[filetype] = filetype_counts.get(filetype, 0) + 1

    for filetype, count in sorted(filetype_counts.items()):
        print(f"  {filetype}: {count} chunks")

    print()
    print("="*80)
    print("✅ INDEX CREATED SUCCESSFULLY")
    print("="*80)
    print(f"\nIndex location: {CHROMA_DIR}")
    print(f"Total indexed chunks: {len(all_documents)}")
    print("\nYou can now use: python3 scripts/smart_query.py")
    print()

if __name__ == "__main__":
    create_index()
